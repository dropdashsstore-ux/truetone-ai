from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime
import requests
import os
import io
from docx import Document
from docx.shared import Inches
import re

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-here'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///truetone.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# API Configuration
CHATANYWHERE_API_KEY = 'sk-c547IW2jbvQv0Q5CurfFWHFbMvFjYUkysAb2DR8NgzDGSM4f'
CHATANYWHERE_API_URL = 'https://api.chatanywhere.tech/v1/chat/completions'

# PayPal Configuration
PAYPAL_CLIENT_ID = 'AUPslxL_eKZd9psSNERgvPHJaWJHFYKWzRqKs-qFx1d81_vLWBF7H4v1YqCsQeMIDyNJ4Xmed_NKdqy9'
PAYPAL_SECRET = 'EH5cyAyKlsUWn-qyWkgXHbLFg-Er374TSwEb-T7p-8SQbsgQhDuTUSYwuJGrsPWTZpn8WZVcrop5RQta'

# Make PayPal client ID available to templates
app.jinja_env.globals.update(PAYPAL_CLIENT_ID=PAYPAL_CLIENT_ID)

# Database Models
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(128))
    plan = db.Column(db.String(20), default='free')  # free, pro, enterprise
    paypal_subscription_id = db.Column(db.String(100))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    def set_password(self, password):
        self.password_hash = generate_password_hash(password)
    
    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

class Usage(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    essay_text = db.Column(db.Text)
    word_count = db.Column(db.Integer)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Routes
@app.route('/')
def home():
    return render_template('home.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        
        # Check if this is the admin account
        if email == 'admintruetone@admin.com' and password == 'adminaitone@2019':
            # Create or get admin user
            admin_user = User.query.filter_by(email=email).first()
            if not admin_user:
                admin_user = User(email=email, plan='enterprise')
                admin_user.set_password(password)
                db.session.add(admin_user)
                db.session.commit()
            else:
                # Ensure admin always has enterprise plan
                admin_user.plan = 'enterprise'
                db.session.commit()
            
            login_user(admin_user)
            return redirect(url_for('humanize'))
        
        # Regular user login
        user = User.query.filter_by(email=email).first()
        
        if user and user.check_password(password):
            login_user(user)
            return redirect(url_for('humanize'))
        else:
            flash('Invalid email or password')
    
    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        
        if User.query.filter_by(email=email).first():
            flash('Email already registered')
            return render_template('register.html')
        
        # Check if this is the admin account
        if email == 'admintruetone@admin.com':
            flash('This email is reserved for admin use. Please use a different email.')
            return render_template('register.html')
        
        user = User(email=email)
        user.set_password(password)
        db.session.add(user)
        db.session.commit()
        
        login_user(user)
        return redirect(url_for('humanize'))
    
    return render_template('register.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('home'))

@app.route('/humanize')
@login_required
def humanize():
    return render_template('humanize.html', user=current_user)

@app.route('/pricing')
def pricing():
    return render_template('pricing.html')

@app.route('/api/humanize', methods=['POST'])
@login_required
def api_humanize():
    data = request.get_json()
    text = data.get('text', '')
    
    # Count words
    word_count = len(text.split())
    
    # Check word limit for free users
    if current_user.plan == 'free' and word_count > 100:
        return jsonify({'error': 'Free plan limited to 100 words. Upgrade to Pro for unlimited words.'}), 400
    
    # Check daily limit for free users
    today = datetime.utcnow().date()
    today_usage = Usage.query.filter(
        Usage.user_id == current_user.id,
        db.func.date(Usage.created_at) == today
    ).count()
    
    if current_user.plan == 'free' and today_usage >= 3:
        return jsonify({'error': 'Free plan limited to 3 essays per day. Upgrade to Pro for unlimited essays.'}), 400
    
    # Call AI API
    try:
        headers = {
            'Authorization': f'Bearer {CHATANYWHERE_API_KEY}',
            'Content-Type': 'application/json'
        }
        
        prompt = f"""Rewrite the following text so it feels 100% human — natural, conversational, and authentic. 
Make it flow like real human writing, with personality, subtle imperfections, and rhythm. 
It must not sound robotic, stiff, or formulaic, and it should be able to pass any AI detector as genuine human writing. 
Preserve the meaning, but make the tone truly human.

Text to humanize:
{text}"""
        
        payload = {
            'model': 'gpt-3.5-turbo',
            'messages': [
                {'role': 'user', 'content': prompt}
            ],
            'max_tokens': 2000,
            'temperature': 0.8
        }
        
        response = requests.post(CHATANYWHERE_API_URL, headers=headers, json=payload)
        response.raise_for_status()
        
        result = response.json()
        humanized_text = result['choices'][0]['message']['content']
        
        # Save usage
        usage = Usage(
            user_id=current_user.id,
            essay_text=text,
            word_count=word_count
        )
        db.session.add(usage)
        db.session.commit()
        
        return jsonify({'humanized_text': humanized_text})
        
    except Exception as e:
        return jsonify({'error': f'Error processing text: {str(e)}'}), 500

@app.route('/api/download/<format>', methods=['POST'])
@login_required
def download_text(format):
    data = request.get_json()
    text = data.get('text', '')
    
    if format == 'txt':
        output = io.BytesIO()
        output.write(text.encode('utf-8'))
        output.seek(0)
        return send_file(
            output,
            as_attachment=True,
            download_name='humanized_text.txt',
            mimetype='text/plain'
        )
    
    elif format == 'docx':
        doc = Document()
        doc.add_heading('Humanized Text', 0)
        doc.add_paragraph(text)
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return send_file(
            output,
            as_attachment=True,
            download_name='humanized_text.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    return jsonify({'error': 'Invalid format'}), 400

@app.route('/api/upgrade-plan', methods=['POST'])
@login_required
def upgrade_plan():
    try:
        data = request.get_json()
        plan = data.get('plan')
        payment_id = data.get('payment_id')
        payer_id = data.get('payer_id')
        email = data.get('email')
        
        if not plan or not payment_id:
            return jsonify({'success': False, 'error': 'Missing required data'}), 400
        
        # Validate plan
        if plan not in ['pro', 'enterprise']:
            return jsonify({'success': False, 'error': 'Invalid plan'}), 400
        
        # Update user's plan
        current_user.plan = plan
        current_user.paypal_subscription_id = payment_id
        db.session.commit()
        
        # Log the upgrade
        print(f"User {current_user.email} upgraded to {plan} plan. Payment ID: {payment_id}")
        
        return jsonify({
            'success': True, 
            'message': 'Plan upgraded successfully',
            'plan': plan
        })
        
    except Exception as e:
        print(f"Error upgrading plan: {str(e)}")
        return jsonify({'success': False, 'error': 'Internal server error'}), 500

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        
        # Create admin account if it doesn't exist
        admin_user = User.query.filter_by(email='admintruetone@admin.com').first()
        if not admin_user:
            admin_user = User(email='admintruetone@admin.com', plan='enterprise')
            admin_user.set_password('adminaitone@2019')
            db.session.add(admin_user)
            db.session.commit()
            print("Admin account created successfully!")
        else:
            # Ensure admin always has enterprise plan
            admin_user.plan = 'enterprise'
            db.session.commit()
            print("Admin account updated with Enterprise Plan!")
    
    app.run(debug=True)
